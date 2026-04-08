#!/usr/bin/env python3
"""
测试公文格式处理工具
创建一个简单的测试文档并演示处理流程
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_test_document():
    """创建一个包含常见格式问题的测试文档"""
    doc = Document()
    
    # 标题（格式有问题：使用了英文括号和逗号）
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("关于召开 2024 年度工作会议的通知 (草稿)")
    run.font.size = Pt(22)
    run.font.bold = True
    
    # 主送机关（格式有问题：使用了英文冒号）
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = recipient.add_run("各部门:")
    run.font.size = Pt(16)
    
    # 正文第一段（缺少首行缩进，使用了英文标点）
    para1 = doc.add_paragraph()
    run = para1.add_run("According to the work arrangement, decided to hold the 2024 annual work conference. Please arrange relevant personnel to attend.")
    run.font.size = Pt(16)
    
    # 一级标题
    heading1 = doc.add_paragraph()
    run = heading1.add_run("一、会议时间")
    run.font.size = Pt(16)
    run.font.bold = True
    
    # 正文（使用了英文数字序号）
    para2 = doc.add_paragraph()
    run = para2.add_run("1. Time: 2024-01-15 9:00 AM")
    run.font.size = Pt(16)
    
    # 二级标题
    heading2 = doc.add_paragraph()
    run = heading2.add_run("（一）会议地点")
    run.font.size = Pt(16)
    
    # 正文
    para3 = doc.add_paragraph()
    run = para3.add_run("会议地点：公司三楼会议室。请提前 10 分钟到场签到。")
    run.font.size = Pt(16)
    
    # 一级标题
    heading3 = doc.add_paragraph()
    run = heading3.add_run("二、参会人员")
    run.font.size = Pt(16)
    
    # 正文（有英文标点）
    para4 = doc.add_paragraph()
    run = para4.add_run("全体中层以上干部参加，特殊情况需要请假者，须经总经理批准 (联系人：张三，电话：123456)。")
    run.font.size = Pt(16)
    
    # 附件
    attachment = doc.add_paragraph()
    run = attachment.add_run("附件：2024 年度工作会议议程")
    run.font.size = Pt(16)
    
    # 落款
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = signature.add_run("公司总经理办公室")
    run.font.size = Pt(16)
    
    # 日期
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date.add_run("2024 年 1 月 10 日")
    run.font.size = Pt(16)
    
    # 保存测试文档
    test_file = "test_input.docx"
    doc.save(test_file)
    print(f"Test document created: {test_file}")
    return test_file


def main():
    """演示完整的处理流程"""
    print("=" * 60)
    print("公文格式处理工具 - 演示流程")
    print("=" * 60)
    print()
    
    # 创建测试文档
    print("Step 0: 创建测试文档...")
    test_file = create_test_document()
    print()
    
    # 诊断
    print("Step 1: 诊断格式问题...")
    print("-" * 60)
    os.system(f"uv run --with python-docx python3 scripts/analyzer.py {test_file}")
    print()
    
    # 修复标点
    print("Step 2: 修复标点符号...")
    print("-" * 60)
    temp_file = "test_temp.docx"
    os.system(f"uv run --with python-docx python3 scripts/punctuation.py {test_file} {temp_file}")
    print()
    
    # 应用公文格式
    print("Step 3: 应用公文格式...")
    print("-" * 60)
    output_file = "test_output.docx"
    os.system(f"uv run --with python-docx python3 scripts/formatter.py {temp_file} {output_file}")
    print()
    
    # 再次诊断
    print("Step 4: 再次诊断确认...")
    print("-" * 60)
    os.system(f"uv run --with python-docx python3 scripts/analyzer.py {output_file}")
    print()
    
    print("=" * 60)
    print("演示完成！")
    print(f"输入文件：{test_file}")
    print(f"输出文件：{output_file}")
    print("=" * 60)


if __name__ == "__main__":
    main()
