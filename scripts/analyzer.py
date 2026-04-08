#!/usr/bin/env python3
"""
公文格式诊断模块 v1.0 - GB/T 9704-2012
严格按照国家行政机关公文格式标准进行诊断
"""

import re
import sys
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 不需要首行缩进的模式（公文标准）
NO_INDENT_PATTERNS = [
    r'^附件 [：:]',                      # 附件：
    r'^附件\d*[：:．.\s]',              # 附件 1：
    r'^附件$',                          # 附件
    r'^主送 [：:]',                      # 主送：
    r'^抄送 [：:]',                      # 抄送：
    r'^印发 [：:]',                      # 印发：
    r'^联系人 [：:]',                    # 联系人：
    r'^（此件公开发布）',              # 公开属性
    r'^（此件依申请公开）',
    r'^（此件不予公开）',
]


def is_no_indent_para(text, alignment):
    """检查是否是不需要首行缩进的段落"""
    # 居中的短文本（主标题）
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    # 特殊行
    for pattern in NO_INDENT_PATTERNS:
        if re.match(pattern, text.strip()):
            return True
    return False


def analyze_punctuation(doc):
    """分析标点符号问题（公文标准）"""
    issues = []
    
    patterns = [
        ('英文括号', r'[\(\)]'),
        ('英文引号', r'["\']'),
        ('英文冒号', r'(?<=[^\d\s]):(?=[^\d/\\])'),
        ('英文逗号', r'(?<=[^\d]),(?=[^\d])'),
        ('英文分号', r';'),
        ('英文问号', r'\?'),
        ('英文叹号', r'!'),
    ]
    
    # 省略号：2 个及以上连续的点（不是省略号格式）
    ellipsis_pattern = r'\.{2,}'
    # 破折号：连续的 -
    dash_pattern = r'--+'
    # 句末英文句号：中文后面的单独句点（不是省略号的一部分）
    period_pattern = r'(?<=[\u4e00-\u9fff])\.(?!\.)'
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        # 只在包含中文的段落中检查
        if not re.search(r'[\u4e00-\u9fff]', text):
            continue
        
        for name, pattern in patterns:
            for match in re.finditer(pattern, text):
                issues.append({
                    'para': i + 1,
                    'type': name,
                    'char': match.group(),
                })
        
        # 检查省略号（连续多个点）
        for match in re.finditer(ellipsis_pattern, text):
            issues.append({'para': i + 1, 'type': '不规范省略号', 'char': match.group()})
        
        # 检查破折号
        for match in re.finditer(dash_pattern, text):
            issues.append({'para': i + 1, 'type': '不规范破折号', 'char': match.group()})
        
        # 检查句末英文句号
        for match in re.finditer(period_pattern, text):
            issues.append({'para': i + 1, 'type': '英文句号', 'char': match.group()})
    
    return issues


def analyze_numbering(doc):
    """分析序号问题（公文标准层次序数）"""
    issues = []
    
    # 公文标准层次序数
    numbering_patterns = {
        '一级标题': r'^[一二三四五六七八九十]+、',           # 一、
        '二级标题': r'^（[一二三四五六七八九十]+）',        # （一）
        '三级标题': r'^\d+\.',                              # 1.
        '四级标题': r'^（\d+）',                            # （1）
        '阿拉伯逗号': r'^\d+,',                             # 1,（不规范）
        '英文括号': r'^\d+\)',                              # 1)（不规范）
    }
    
    found_styles = defaultdict(list)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        for style_name, pattern in numbering_patterns.items():
            if re.match(pattern, text):
                found_styles[style_name].append(i + 1)
                break
    
    # 检查阿拉伯数字序号是否规范（应该用 1. 而不是 1,或 1)）
    if found_styles['阿拉伯逗号'] or found_styles['英文括号']:
        issues.append({
            'type': '序号格式不规范',
            'detail': '第三层应该用"1."而不是"1,"或"1)"',
        })
    
    # 检查阿拉伯数字序号风格是否统一
    if found_styles['三级标题'] and (found_styles['阿拉伯逗号'] or found_styles['英文括号']):
        issues.append({
            'type': '序号格式不统一',
            'detail': f"同时存在规范和不规范的序号格式",
        })
    
    # 检查层级跳跃
    has_level1 = bool(found_styles['一级标题'])
    has_level2 = bool(found_styles['二级标题'])
    has_level3 = bool(found_styles['三级标题'])
    has_level4 = bool(found_styles['四级标题'])
    
    # 如果同时存在多级标题，检查层级顺序
    levels = [has_level1, has_level2, has_level3, has_level4]
    if sum(levels) > 1:
        # 检查是否有跳跃（如从一级直接到三级）
        for i in range(len(levels) - 2):
            if levels[i] and i + 2 < len(levels) and levels[i + 2] and not levels[i + 1]:
                issues.append({
                    'type': '层级跳跃',
                    'detail': f'从第{i+1}层直接跳到第{i+3}层，缺少第{i+2}层',
                })
    
    return issues


def analyze_paragraph_format(doc):
    """分析段落格式问题（公文标准）"""
    issues = []
    
    indent_issues = []
    line_spacing_values = defaultdict(list)
    alignment_issues = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 跳过空段落和短段落（可能是标题）
        if not text or len(text) < 10:
            continue
        
        alignment = para.paragraph_format.alignment
        
        # 跳过不需要缩进的段落
        if is_no_indent_para(text, alignment):
            continue
        
        pf = para.paragraph_format
        
        # 检查首行缩进
        indent = pf.first_line_indent
        if indent is None or indent == Pt(0) or (hasattr(indent, 'pt') and indent.pt == 0):
            indent_issues.append(i + 1)
        
        # 记录行距
        if pf.line_spacing is not None:
            line_spacing_values[str(pf.line_spacing)].append(i + 1)
        
        # 检查对齐方式（正文应该两端对齐）
        if alignment == WD_ALIGN_PARAGRAPH.LEFT and len(text) > 30:
            # 长段落左对齐可能不规范
            alignment_issues.append(i + 1)
    
    if indent_issues:
        issues.append({
            'type': '缺少首行缩进',
            'paras': indent_issues
        })
    
    if len(line_spacing_values) > 1:
        issues.append({
            'type': '行距不统一',
            'detail': f"存在 {len(line_spacing_values)} 种不同行距",
        })
    
    if alignment_issues:
        issues.append({
            'type': '对齐方式不规范',
            'detail': f'第{", ".join(map(str, alignment_issues[:5]))}段等建议改为两端对齐',
        })
    
    return issues


def analyze_font(doc):
    """分析字体问题（公文标准）"""
    issues = []
    
    font_names = set()
    font_sizes = set()
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        for run in para.runs:
            if run.font.name:
                font_names.add(run.font.name)
            if run.font.size:
                font_sizes.add(str(run.font.size))
    
    # 公文一般不超过 3 种字体
    if len(font_names) > 4:
        issues.append({
            'type': '字体种类过多',
            'detail': f"检测到 {len(font_names)} 种字体混用，公文一般不超过 3 种"
        })
    
    # 公文正文一般只有一种字号（三号）
    if len(font_sizes) > 3:
        issues.append({
            'type': '字号不统一',
            'detail': f"检测到 {len(font_sizes)} 种字号混用，公文正文应统一为三号"
        })
    
    return issues


def analyze_structure(doc):
    """分析公文结构完整性"""
    issues = []
    
    has_title = False
    has_recipient = False
    has_body = False
    has_signature = False
    has_date = False
    
    date_pattern = r'^\d{4}年\d{1,2}月\d{1,2}日$'
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        alignment = para.paragraph_format.alignment
        
        # 检测标题
        if i < 5 and (
            re.match(r'^关于.+的 (通知 | 报告 | 请示 | 函 | 意见 | 决定 | 公告 | 通报 | 批复)', text) or
            (15 < len(text) < 80 and alignment == WD_ALIGN_PARAGRAPH.CENTER)
        ):
            has_title = True
        
        # 检测主送机关
        if re.match(r'^[\u4e00-\u9fff]+[：:]$', text) and len(text) < 20:
            has_recipient = True
        
        # 检测正文
        if len(text) > 50:
            has_body = True
        
        # 检测落款
        if re.search(r'(公司 | 局 | 委 | 部 | 厅 | 院 | 所 | 中心 | 办公室 | 集团 | 银行 | 学校 | 大学 | 医院)$', text):
            has_signature = True
        
        # 检测日期
        if re.match(date_pattern, text):
            has_date = True
    
    # 结构完整性检查
    if not has_title:
        issues.append({
            'type': '缺少公文标题',
            'detail': '建议添加"关于 XXX 的通知/报告"等标题',
        })
    
    if not has_body:
        issues.append({
            'type': '缺少正文',
            'detail': '未检测到足够的正文内容',
        })
    
    # 日期和落款通常成对出现
    if has_signature and not has_date:
        issues.append({
            'type': '缺少成文日期',
            'detail': '有落款但未检测到成文日期',
        })
    
    return issues


def print_report(results):
    """打印诊断报告（公文风格）"""
    print('=' * 50)
    print('           公文格式诊断报告')
    print('=' * 50)
    print()
    
    total = 0
    
    # 标点问题
    punct = results['punctuation']
    if punct:
        by_type = defaultdict(list)
        for issue in punct:
            by_type[issue['type']].append(issue['para'])
        
        print(f"【标点问题】共 {len(punct)} 处")
        for issue_type, paras in by_type.items():
            unique_paras = sorted(set(paras))
            if len(unique_paras) > 5:
                para_str = f"第{unique_paras[0]}、{unique_paras[1]}...{unique_paras[-1]}段"
            else:
                para_str = f"第{', '.join(map(str, unique_paras))}段"
            print(f"  - {issue_type}: {para_str}")
        print()
        total += len(punct)
    
    # 序号问题
    num = results['numbering']
    if num:
        print(f"【序号问题】共 {len(num)} 处")
        for issue in num:
            print(f"  - {issue['type']}: {issue.get('detail', '')}")
        print()
        total += len(num)
    
    # 段落问题
    para = results['paragraph']
    if para:
        print(f"【段落格式问题】共 {len(para)} 处")
        for issue in para:
            if issue['type'] == '缺少首行缩进':
                paras = issue['paras']
                if len(paras) > 5:
                    para_str = f"第{paras[0]}、{paras[1]}...等{len(paras)}段"
                else:
                    para_str = f"第{', '.join(map(str, paras))}段"
                print(f"  - {issue['type']}: {para_str}")
            else:
                print(f"  - {issue['type']}: {issue.get('detail', '')}")
        print()
        total += len(para)
    
    # 字体问题
    font = results['font']
    if font:
        print(f"【字体问题】共 {len(font)} 处")
        for issue in font:
            print(f"  - {issue['type']}: {issue.get('detail', '')}")
        print()
        total += len(font)
    
    # 结构问题
    struct = results['structure']
    if struct:
        print(f"【公文结构问题】共 {len(struct)} 处")
        for issue in struct:
            print(f"  - {issue['type']}: {issue.get('detail', '')}")
        print()
        total += len(struct)
    
    # 总结
    print('-' * 50)
    if total == 0:
        print('✓ 未发现明显格式问题，符合公文格式要求')
    else:
        print(f'共发现 {total} 处格式问题')
        print()
        print('建议：')
        if results['punctuation']:
            print('  - 运行 punctuation.py 修复标点问题')
        if results['paragraph'] or results['font']:
            print('  - 运行 formatter.py 统一公文格式')
    print()


def main():
    if len(sys.argv) < 2:
        print('Usage: python analyzer.py input.docx [--json]')
        sys.exit(1)
    
    input_file = sys.argv[1]
    print(f'Analyzing: {input_file}')
    print()
    
    doc = Document(input_file)
    
    results = {
        'punctuation': analyze_punctuation(doc),
        'numbering': analyze_numbering(doc),
        'paragraph': analyze_paragraph_format(doc),
        'font': analyze_font(doc),
        'structure': analyze_structure(doc),
    }
    
    if '--json' in sys.argv:
        import json
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        print_report(results)


if __name__ == '__main__':
    main()
