#!/usr/bin/env python3
"""
文档格式统一 v1.0 - GB/T 9704-2012 党政机关公文格式
严格按照国家标准设置页面、字体、字号、行距等

公文标准：
- 页边距：上 37mm，下 35mm，左 28mm，右 26mm
- 主标题：居中，二号（22pt），方正小标宋简体
- 主送机关：顶格，三号仿宋
- 正文：三号仿宋 GB2312，首行缩进 2 字符，行距 28 磅
- 一级标题："一、" 三号黑体，首行缩进 2 字符
- 二级标题："（一）" 三号楷体 GB2312，首行缩进 2 字符
- 三级标题："1." 三号仿宋 GB2312，首行缩进 2 字符
- 四级标题："（1）" 三号仿宋 GB2312，首行缩进 2 字符
- 落款：右对齐，三号仿宋
- 附件：首行缩进 2 字符，三号仿宋
- 页码：四号宋体，一字线，奇右偶左，距版心下边缘 7mm
"""

import sys
import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 字号对照：二号=22pt，三号=16pt，四号=14pt，小四=12pt
# 2 字符缩进 = 2 × 16pt = 32pt（三号字）

# 预设配置目录
PRESETS_DIR = Path(__file__).parent.parent / "presets"


def load_preset(preset_name):
    """加载预设配置"""
    preset_file = PRESETS_DIR / f"{preset_name}.json"
    if preset_file.exists():
        try:
            with open(preset_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f'Warning: Failed to load preset {preset_name}: {e}')
    return None


# 标准公文格式预设（GB/T 9704-2012）
OFFICIAL_PRESET = {
    'name': '党政机关公文格式（GB/T 9704-2012）',
    'page': {
        'top': 3.7,      # 上边距 37mm
        'bottom': 3.5,   # 下边距 35mm
        'left': 2.8,     # 左边距 28mm
        'right': 2.6,    # 右边距 26mm
    },
    # 主标题：二号方正小标宋简体，居中
    'title': {
        'font_cn': '方正小标宋简体',
        'font_en': 'Times New Roman',
        'size': 22,      # 二号
        'bold': False,
        'align': 'center',
        'indent': 0,
        'line_spacing': 28,
    },
    # 主送机关：三号仿宋，顶格
    'recipient': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,      # 三号
        'bold': False,
        'align': 'left',
        'indent': 0,     # 顶格
        'line_spacing': 28,
    },
    # 一级标题：三号黑体，"一、"，首行缩进 2 字符
    'heading1': {
        'font_cn': '黑体',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'left',
        'indent': 32,    # 2 字符缩进
        'line_spacing': 28,
    },
    # 二级标题：三号楷体 GB2312，"（一）"，首行缩进 2 字符
    'heading2': {
        'font_cn': '楷体_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'left',
        'indent': 32,
        'line_spacing': 28,
    },
    # 三级标题：三号仿宋 GB2312，"1."，首行缩进 2 字符
    'heading3': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'left',
        'indent': 32,
        'line_spacing': 28,
    },
    # 四级标题：三号仿宋 GB2312，"（1）"，首行缩进 2 字符
    'heading4': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'left',
        'indent': 32,
        'line_spacing': 28,
    },
    # 正文：三号仿宋 GB2312，首行缩进 2 字符，行距 28 磅
    'body': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'justify',
        'indent': 32,    # 2 字符 = 2×16pt
        'line_spacing': 28,
    },
    # 落款单位：三号仿宋，右对齐
    'signature': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'right',
        'indent': 0,
        'line_spacing': 28,
    },
    # 成文日期：三号仿宋，右对齐
    'date': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'right',
        'indent': 0,
        'line_spacing': 28,
    },
    # 附件说明：三号仿宋，首行缩进 2 字符
    'attachment': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'justify',
        'indent': 32,
        'line_spacing': 28,
    },
    # 结束语（特此通知等）：三号仿宋，首行缩进
    'closing': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'left',
        'indent': 32,
        'line_spacing': 28,
    },
    # 附注：三号仿宋，居左空 2 字
    'note': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'left',
        'indent': 32,
        'line_spacing': 28,
    },
    # 抄送单位：三号仿宋，左右各空 1 字
    'copy_to': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 16,
        'bold': False,
        'align': 'justify',
        'indent': 16,    # 1 字符
        'line_spacing': 28,
    },
    # 表格内容：小四号仿宋
    'table': {
        'font_cn': '仿宋_GB2312',
        'font_en': 'Times New Roman',
        'size': 14,      # 小四
        'bold': False,
        'line_spacing': 20,
        'header_bold': True,
        'first_line_indent': 0,
    },
    # 页码设置
    'page_number': {
        'enabled': True,
        'font': '宋体',
        'size': 14,      # 四号
    },
}


def remove_background(doc):
    """移除页面背景颜色"""
    body = doc._body._body
    document = body.getparent()
    for elem in list(document):
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == 'background':
            document.remove(elem)
    
    for para in doc.paragraphs:
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        for run in para.runs:
            run.font.highlight_color = None
            rPr = run._r.get_or_add_rPr()
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                rPr.remove(shd)


def _iter_block_items(doc):
    """Yield paragraphs and tables in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, doc)
        elif child.tag.endswith('}tbl'):
            yield Table(child, doc)


def _set_table_borders(table, size_pt=0.5, color="000000"):
    """设置表格边框"""
    size = max(1, int(size_pt * 8))  # OOXML border size is in 1/8 pt
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def _set_table_cell_margins(table, top_cm=0.0, bottom_cm=0.0, left_cm=0.05, right_cm=0.05):
    """设置表格单元格边距"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    cell_mar = tbl_pr.find(qn('w:tblCellMar'))
    if cell_mar is None:
        cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(cell_mar)

    def _set_side(tag, cm_value):
        node = cell_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            cell_mar.append(node)
        node.set(qn('w:type'), 'dxa')
        node.set(qn('w:w'), str(int(Cm(cm_value).twips)))

    _set_side('top', top_cm)
    _set_side('bottom', bottom_cm)
    _set_side('left', left_cm)
    _set_side('right', right_cm)


def _set_table_width_percent(table, percent=100):
    """设置表格宽度百分比"""
    percent = max(1, min(100, int(percent)))
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'pct')
    tbl_w.set(qn('w:w'), str(percent * 50))  # 50ths of a percent


def _set_table_indent(table, indent_twips=0):
    """设置表格缩进"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_ind.set(qn('w:w'), str(int(indent_twips)))


def _text_weight(text):
    """计算文本权重（用于列宽分配）"""
    weight = 0.0
    for ch in text:
        if ord(ch) < 128:
            weight += 0.5
        else:
            weight += 1.0
    return weight


def _normalize_pcts(weights, min_pct, max_pct):
    """归一化百分比"""
    total = sum(weights) or 1.0
    pcts = [w / total * 100 for w in weights]

    # Clamp low
    for i, v in enumerate(pcts):
        if v < min_pct:
            pcts[i] = min_pct
    # Clamp high
    for i, v in enumerate(pcts):
        if v > max_pct:
            pcts[i] = max_pct

    # Renormalize to 100
    total = sum(pcts) or 1.0
    return [v / total * 100 for v in pcts]


def _set_table_col_widths_by_content(table, min_pct=8, max_pct=45):
    """根据内容自动设置列宽"""
    if not table.rows:
        return
    col_count = max(len(row.cells) for row in table.rows)
    if col_count == 0:
        return

    max_weights = [1.0] * col_count
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            text = ''.join(p.text for p in cell.paragraphs).strip()
            if text:
                max_weights[c_idx] = max(max_weights[c_idx], _text_weight(text))

    pcts = _normalize_pcts(max_weights, min_pct, max_pct)

    # Set table grid + cell widths in pct
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement('w:tblGrid')
        tbl.insert(0, tbl_grid)
    else:
        for child in list(tbl_grid):
            tbl_grid.remove(child)

    for pct in pcts:
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), str(int(pct * 50)))  # pct in 1/50th %
        tbl_grid.append(grid_col)

    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tc_pr = tc.tcPr
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:type'), 'pct')
            tc_w.set(qn('w:w'), str(int(pcts[c_idx] * 50)))


def _insert_paragraph_after_table(table, text=""):
    """在表格后插入段落"""
    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_before_table(table, text=""):
    """在表格前插入段落"""
    p = OxmlElement("w:p")
    table._tbl.addprevious(p)
    para = Paragraph(p, table._parent)
    if text:
        para.add_run(text)
    return para


def _insert_paragraph_after_paragraph(paragraph, text=""):
    """在段落后插入段落"""
    p = OxmlElement("w:p")
    paragraph._p.addnext(p)
    para = Paragraph(p, paragraph._parent)
    if text:
        para.add_run(text)
    return para


def _is_numeric_text(text):
    """检查是否为数字文本"""
    text = text.replace(',', '').replace('％', '%').strip()
    if not text:
        return False
    return re.match(r'^[-+]?\d+(?:\.\d+)?%?$', text) is not None


def _is_short_text(text, max_len=4):
    """检查是否为短文本"""
    text = text.strip()
    return 0 < len(text) <= max_len


def _is_table_title(text):
    """检查是否为表格标题"""
    text = text.strip()
    if not text:
        return False
    if len(text) > 30:
        return False
    return re.match(r'^表\s*(?:\d+|[一二三四五六七八九十]+)(?:[\-\—\._、]\d+)?', text) is not None


def _is_table_unit(text):
    """检查是否为表格单位"""
    text = text.strip()
    if not text:
        return False
    if len(text) > 20:
        return False
    return re.match(r'^单位\s*[:：]', text) is not None


def _set_cell_borders(cell, size_pt=0.5, color="000000"):
    """设置单元格边框"""
    size = max(1, int(size_pt * 8))
    tc = cell._tc
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)

    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), color)
        borders.append(elem)


def detect_para_type(text, index, total, alignment, all_texts):
    """
    检测段落类型（公文标准）
    
    返回：'title', 'recipient', 'heading1', 'heading2', 'heading3', 'heading4', 
          'body', 'signature', 'date', 'attachment', 'closing', 'note', 'copy_to'
    """
    text = text.strip()
    if not text:
        return 'empty'
    
    # ===== 一级标题："一、" "二、" 等 =====
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 'heading1'
    
    # ===== 二级标题："（一）" "（二）" 等 =====
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 'heading2'
    if re.match(r'^\([一二三四五六七八九十]+\)', text):
        return 'heading2'
    
    # ===== 三级标题："1." "2." 等 =====
    if re.match(r'^\d+\.\s*\S', text) and len(text) < 60:
        return 'heading3'
    
    # ===== 四级标题："（1）" "（2）" 等 =====
    if re.match(r'^（\d+）', text) and len(text) < 60:
        return 'heading4'
    if re.match(r'^\(\d+\)', text) and len(text) < 60:
        return 'heading4'
    
    # ===== 主送机关：XXX：或 XXX: =====
    if re.match(r'^[\u4e00-\u9fff]+[：:]$', text) and len(text) < 20:
        return 'recipient'
    
    # ===== 附件行 =====
    if re.match(r'^附件 [：:]\s*', text):
        return 'attachment'
    if re.match(r'^附件\d*[：:．.\s]', text):
        return 'attachment'
    if re.match(r'^附件$', text):
        return 'attachment'
    
    # ===== 结束语 =====
    closing_patterns = [
        r'^特此 (说明 | 通知 | 报告 | 函复 | 函告 | 批复 | 公告 | 通报) 。?$',
        r'^此致$',
        r'^敬礼 [！!]?$',
        r'^以上 (报告 | 意见 | 方案).{0,10}$',
        r'^妥否.{0,10}$',
        r'^请.{0,15}(批示 | 审批 | 审议 | 指示 | 核准)。?$',
    ]
    for pattern in closing_patterns:
        if re.match(pattern, text):
            return 'closing'
    
    # ===== 成文日期 =====
    date_patterns = [
        r'^\d{4}年\d{1,2}月\d{1,2}日$',
        r'^\d{4}\.\d{1,2}\.\d{1,2}$',
        r'^\d{4}/\d{1,2}/\d{1,2}$',
        r'^\d{4}-\d{1,2}-\d{1,2}$',
    ]
    for pattern in date_patterns:
        if re.match(pattern, text):
            return 'date'
    
    # ===== 附注 =====
    if re.match(r'^（此件', text) or re.match(r'^\(此件', text):
        return 'note'
    
    # ===== 抄送单位 =====
    if re.match(r'^抄送 [：:]', text):
        return 'copy_to'
    if re.match(r'^分送 [：:]', text):
        return 'copy_to'
    
    # ===== 落款单位 =====
    if index >= total - 10 and len(text) < 30:
        if re.search(r'(公司 | 局 | 委 | 部 | 厅 | 院 | 所 | 中心 | 办公室 | 集团 | 银行 | 学校 | 大学 | 医院)$', text):
            return 'signature'
        # 检查下文是否有日期
        remaining_texts = all_texts[all_texts.index(text)+1:] if text in all_texts else []
        for next_text in remaining_texts[:3]:
            for pattern in date_patterns:
                if re.match(pattern, next_text.strip()):
                    return 'signature'
    
    # ===== 主标题 =====
    if index < 5:
        title_patterns = [
            r'^关于.+的 (通知 | 报告 | 请示 | 函 | 意见 | 决定 | 公告 | 通报 | 批复 | 说明 | 方案 | 总结 | 汇报 | 复函 | 答复 | 建议)$',
            r'^.{2,30}(通知 | 报告 | 请示 | 函 | 意见 | 决定 | 公告 | 通报 | 批复 | 工作方案 | 工作总结 | 实施方案 | 管理办法 | 暂行规定)$',
        ]
        for pattern in title_patterns:
            if re.match(pattern, text):
                return 'title'
        
        # 较长的标题（20-80 字符），不以标点结尾
        if 15 < len(text) < 80 and not re.search(r'[。！？，、；：]$', text):
            if not re.match(r'^[一二三四五六七八九十\d（(]', text):
                return 'title'
        
        # 居中的短文本
        if alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 60:
            return 'title'
    
    # ===== 其他都是正文 =====
    return 'body'


def _split_heading_by_punct(paragraph):
    """将"标题 + 标点 + 正文"拆分为标题段 + 正文段"""
    text = paragraph.text.strip()
    if not text:
        return False

    if not (
        re.match(r'^[一二三四五六七八九十]+、', text) or
        re.match(r'^（[一二三四五六七八九十]+）', text) or
        re.match(r'^\([一二三四五六七八九十]+\)', text) or
        re.match(r'^\d+\.\s*\S', text) or
        re.match(r'^（\d+）', text) or
        re.match(r'^\(\d+\)', text)
    ):
        return False

    punct_positions = []
    for ch in ('：', ':', '。'):
        pos = text.find(ch)
        if pos != -1:
            punct_positions.append(pos)
    if not punct_positions:
        return False
    
    split_idx = min(punct_positions)
    head = text[:split_idx + 1].strip()
    tail = text[split_idx + 1:].strip()
    if not tail:
        return False

    paragraph.text = head
    new_para = _insert_paragraph_after_paragraph(paragraph, text=tail)
    return new_para is not None


def set_font(run, font_cn, font_en, size, bold=False):
    """
    设置字体，同时清除原有格式（斜体、下划线、颜色）
    """
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = False
    run.font.underline = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.strike = False
    run.font.double_strike = False
    run.font.subscript = False
    run.font.superscript = False
    
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:cs'), font_en)


def format_paragraph(para, fmt, para_type):
    """格式化段落"""
    pf = para.paragraph_format
    
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(fmt.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    
    indent = fmt.get('indent', 0)
    if indent > 0:
        pf.first_line_indent = Pt(indent)
    else:
        pf.first_line_indent = Pt(0)
    
    ls = fmt.get('line_spacing', 28)
    if ls:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls)
    else:
        pf.line_spacing = 1.5
    
    pf.space_before = Pt(fmt.get('space_before', 0))
    pf.space_after = Pt(fmt.get('space_after', 0))
    
    for run in para.runs:
        set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))


def add_page_number(doc, font_name="宋体", size=14):
    """
    添加页码（公文标准）
    - 四号宋体
    - 一字线格式：— 1 —
    - 奇数页居右，偶数页居左
    - 距版心下边缘约 7mm
    """
    try:
        doc.settings.odd_and_even_pages_header_footer = True
    except Exception:
        settings_el = doc.settings._element
        if settings_el.find(qn('w:evenAndOddHeaders')) is None:
            settings_el.append(OxmlElement('w:evenAndOddHeaders'))

    for section in doc.sections:
        section.odd_and_even_pages_header_footer = True
        section.footer_distance = Cm(0.7)  # 7mm

        odd_footer = section.footer
        even_footer = section.even_page_footer
        odd_footer.is_linked_to_previous = False
        even_footer.is_linked_to_previous = False

        for f in (odd_footer, even_footer):
            for para in f.paragraphs:
                para.clear()

        def _build_footer_line(footer, align, pad_fullwidth):
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.alignment = align

            if pad_fullwidth:
                run0 = para.add_run(" ")
                set_font(run0, font_name, font_name, size, bold=False)

            run1 = para.add_run("— ")
            set_font(run1, font_name, font_name, size, bold=False)

            run2 = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run2._r.append(fldChar1)
            set_font(run2, font_name, font_name, size, bold=False)

            run3 = para.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run3._r.append(instrText)
            set_font(run3, font_name, font_name, size, bold=False)

            run4 = para.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run4._r.append(fldChar2)
            set_font(run4, font_name, font_name, size, bold=False)

            run5 = para.add_run(" —")
            set_font(run5, font_name, font_name, size, bold=False)

            if not pad_fullwidth:
                run6 = para.add_run(" ")
                set_font(run6, font_name, font_name, size, bold=False)

        _build_footer_line(odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, pad_fullwidth=True)
        _build_footer_line(even_footer, WD_ALIGN_PARAGRAPH.LEFT, pad_fullwidth=False)


def format_document(input_path, output_path, preset=None):
    """格式化文档（公文标准）"""
    if preset is None:
        preset = OFFICIAL_PRESET
    elif isinstance(preset, str):
        if preset == 'custom':
            loaded = load_preset('custom')
            if loaded:
                preset = loaded
            else:
                print('Custom preset not found, using official preset')
                preset = OFFICIAL_PRESET
        else:
            loaded = load_preset(preset)
            if loaded:
                preset = loaded
            else:
                print(f'Preset {preset} not found, using official preset')
                preset = OFFICIAL_PRESET
    
    print(f'Preset: {preset.get("name", "党政机关公文格式（GB/T 9704-2012）")}')
    print(f'Input: {input_path}')
    
    doc = Document(input_path)

    # 拆分标题和正文
    for para in list(doc.paragraphs):
        _split_heading_by_punct(para)

    total_paras = len(doc.paragraphs)
    all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # 1. 移除背景
    print('1. Removing background...')
    remove_background(doc)
    
    # 2. 设置页面边距
    print('2. Setting page margins...')
    page = preset['page']
    for section in doc.sections:
        section.top_margin = Cm(page['top'])
        section.bottom_margin = Cm(page['bottom'])
        section.left_margin = Cm(page['left'])
        section.right_margin = Cm(page['right'])
    
    # 3. 格式化段落
    print('3. Formatting paragraphs...')
    stats = {}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        para_type = detect_para_type(
            text, i, total_paras, 
            para.paragraph_format.alignment,
            all_texts
        )
        
        fmt_key = para_type if para_type in preset else 'body'
        fmt = preset.get(fmt_key, preset['body'])
        
        format_paragraph(para, fmt, para_type)
        stats[para_type] = stats.get(para_type, 0) + 1
        
        preview = text[:35] + '...' if len(text) > 35 else text
        print(f'   [{para_type:12}] {preview}')
    
    # 4. 处理表格
    print('4. Formatting tables...')
    body_fmt = preset.get('body', {})
    table_fmt = preset.get('table', {})
    
    tbl_font_cn = table_fmt.get('font_cn', body_fmt.get('font_cn', '仿宋_GB2312'))
    tbl_font_en = table_fmt.get('font_en', body_fmt.get('font_en', 'Times New Roman'))
    tbl_size = table_fmt.get('size', body_fmt.get('size', 16))
    tbl_line_spacing = table_fmt.get('line_spacing', 20)
    tbl_header_bold = table_fmt.get('header_bold', True)
    tbl_first_line_indent = table_fmt.get('first_line_indent', 0)

    blocks = list(_iter_block_items(doc))
    for idx, block in enumerate(blocks):
        if not isinstance(block, Table):
            continue

        table = block
        
        # 表格优化
        table.autofit = False
        _set_table_width_percent(table, 100)
        _set_table_indent(table, 0)
        _set_table_borders(table, size_pt=0.5)
        _set_table_cell_margins(table, top_cm=0.0, bottom_cm=0.0, left_cm=0.05, right_cm=0.05)
        _set_table_col_widths_by_content(table, min_pct=8, max_pct=45)

        # 表格前空一行
        prev_block = blocks[idx - 1] if idx - 1 >= 0 else None
        if isinstance(prev_block, Paragraph):
            if prev_block.text.strip():
                if _is_table_title(prev_block.text) or _is_table_unit(prev_block.text):
                    _insert_paragraph_before_paragraph(prev_block, text="")
                else:
                    _insert_paragraph_before_table(table, text="")
        elif isinstance(prev_block, Table):
            _insert_paragraph_after_table(prev_block, text="")
        elif idx == 0:
            _insert_paragraph_before_table(table, text="")

        # 表题居中
        if prev_block and isinstance(prev_block, Paragraph) and _is_table_title(prev_block.text):
            prev_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev_block.paragraph_format.space_before = Pt(0)
            prev_block.paragraph_format.space_after = Pt(0)
            prev_block.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 表注右对齐
        next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
        unit_para = None
        if isinstance(next_block, Paragraph) and _is_table_unit(next_block.text):
            unit_para = next_block
            unit_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            unit_para.paragraph_format.space_before = Pt(tbl_size * 0.5)
            unit_para.paragraph_format.space_after = Pt(0)
            unit_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 表格内容
        serial_col_idx = None
        if table.rows:
            header_cells = table.rows[0].cells
            for c_idx, cell in enumerate(header_cells):
                head_text = ''.join(p.text for p in cell.paragraphs).strip()
                if '序号' in head_text or head_text == '序':
                    serial_col_idx = c_idx
                    break

        for row_idx, row in enumerate(table.rows):
            row.height = Cm(0.7)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            for col_idx, cell in enumerate(row.cells):
                _set_cell_borders(cell, size_pt=0.5)

                cell_text = ''.join(p.text for p in cell.paragraphs).strip()
                for para in cell.paragraphs:
                    if para.text.strip():
                        is_header = (row_idx == 0 and tbl_header_bold)
                        for run in para.runs:
                            set_font(run, tbl_font_cn, tbl_font_en, tbl_size, bold=(tbl_header_bold and is_header))

                    para.paragraph_format.first_line_indent = Pt(tbl_first_line_indent)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    para.paragraph_format.line_spacing = Pt(tbl_line_spacing)

                    # 对齐策略
                    if row_idx == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif '合计' in cell_text or '总计' in cell_text:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif serial_col_idx is not None and col_idx == serial_col_idx:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif _is_numeric_text(cell_text):
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif _is_short_text(cell_text, 4):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 表格后空一行
        if unit_para is not None:
            after_unit = blocks[idx + 2] if idx + 2 < len(blocks) else None
            if not (isinstance(after_unit, Paragraph) and not after_unit.text.strip()):
                _insert_paragraph_after_paragraph(unit_para, text="")
        else:
            next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
            if not (isinstance(next_block, Paragraph) and not next_block.text.strip()):
                _insert_paragraph_after_table(table, text="")
    
    # 5. 添加页码
    page_number_cfg = preset.get('page_number', {})
    if page_number_cfg.get('enabled', True):
        print('5. Adding page numbers...')
        pn_font = page_number_cfg.get('font', '宋体')
        pn_size = page_number_cfg.get('size', 14)
        add_page_number(doc, font_name=pn_font, size=pn_size)
    else:
        print('5. Skipping page numbers...')
    
    # 保存
    doc.save(output_path)
    
    print()
    print('=' * 50)
    print('Statistics:')
    for k, v in sorted(stats.items()):
        if v > 0:
            print(f'  {k}: {v}')
    print(f'Output: {output_path}')


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='公文格式排版工具（GB/T 9704-2012）')
    parser.add_argument('input', help='输入文件路径')
    parser.add_argument('output', help='输出文件路径')
    parser.add_argument('--preset', default='official', 
                       help='预设配置：official（默认）| custom | 其他自定义配置')
    
    args = parser.parse_args()
    
    format_document(args.input, args.output, preset=args.preset)
